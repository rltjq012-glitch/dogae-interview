import sys
import os
import locale

# -------------------------------------------------------------------------
# [0] 인코딩/로케일 하드닝
# -------------------------------------------------------------------------
# 일부 서버/컨테이너 환경(특히 LANG/LC_ALL이 설정되지 않은 리눅스 컨테이너)은
# 기본 로케일이 'C'(ASCII)로 잡혀 있어, 한글이 포함된 긴 프롬프트를 전송할 때
# 내부 라이브러리(httpx 등)가 'ascii' codec으로 인코딩을 시도하다가
# "'ascii' codec can't encode characters..." 오류를 일으킬 수 있습니다.
# 앱 시작 시점에 가능한 로케일들을 순서대로 시도해 UTF-8 환경을 강제합니다.
for _loc in ("C.UTF-8", "en_US.UTF-8", "ko_KR.UTF-8", "Korean_Korea.65001", ""):
    try:
        locale.setlocale(locale.LC_ALL, _loc)
        break
    except Exception:
        continue
os.environ["PYTHONUTF8"] = "1"
os.environ.setdefault("PYTHONIOENCODING", "utf-8")
os.environ.setdefault("LANG", "C.UTF-8")
os.environ.setdefault("LC_ALL", "C.UTF-8")
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        try:
            _stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

import streamlit as st
import time
import pymupdf
import re
from google import genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
st.set_page_config(page_title="도개고 면접 마스터", layout="wide")
# -------------------------------------------------------------------------
# [1] 스마트 PDF 및 제미나이 통신 함수
# -------------------------------------------------------------------------
def extract_text_from_pdf(uploaded_file):
    doc = pymupdf.open(stream=uploaded_file.read(), filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text
def call_gemini(prompt, api_key):
    # API 키에 복사/붙여넣기 과정에서 섞여 들어간 공백/개행 문자를 제거합니다.
    # (개행이 섞인 키가 요청 헤더/URL 구성에 사용되면 인코딩 오류의 원인이 될 수 있습니다.)
    api_key = (api_key or "").strip()
    # 프롬프트에 서로게이트(짝이 맞지 않는 유니코드) 문자가 섞여 있으면 인코딩 단계에서
    # 오류가 날 수 있으므로, UTF-8로 안전하게 정규화해 둡니다.
    prompt = prompt.encode("utf-8", errors="ignore").decode("utf-8")

    client = genai.Client(api_key=api_key)

    available_models = []
    try:
        for m in client.models.list():
            name = m.name.replace("models/", "") if m.name.startswith("models/") else m.name
            if "gemini-1.5" in name or "gemini-flash" in name:
                available_models.append(name)
    except:
        available_models = ["gemini-1.5-flash", "gemini-1.5-pro"]

    if not available_models:
        available_models = ["gemini-1.5-flash-latest"]

    models_to_try = sorted(available_models, key=lambda x: "flash" not in x)

    last_error = ""
    for target_model in models_to_try[:3]:
        try:
            response = client.models.generate_content(model=target_model, contents=prompt)
            return response.text
        except UnicodeEncodeError as e:
            # 서버/컨테이너의 로케일이 ASCII로 잡혀 있어 한글 프롬프트를 인코딩하지 못하는 경우.
            # 파일 상단의 로케일 하드닝으로 대부분 예방되지만, 혹시 남아 있다면 진단 메시지를 남깁니다.
            last_error = f"UnicodeEncodeError: {e} (서버 로케일이 UTF-8이 아닐 수 있습니다. PYTHONUTF8=1 환경변수로 앱을 재실행해 보세요.)"
            time.sleep(2)
        except Exception as e:
            last_error = f"{type(e).__name__}: {e}"
            time.sleep(2)

    raise Exception(f"AI 모델 통신 실패 (마지막 에러: {last_error})")
# -------------------------------------------------------------------------
# [2] 워드 표 생성 및 가독성/여백 최적화 엔진
# -------------------------------------------------------------------------
def set_document_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)
def add_parsed_text_to_cell(cell, text):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 != 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
def add_intro_paragraphs(doc, text):
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3

        if line.startswith("### 📄"):
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
        else:
            parts = line.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0: run.bold = True
def add_passage_block_to_doc(doc, block_text):
    """### 📄 로 시작하는 '제시문' 블록을 워드 문서에 렌더링합니다 (학생/교사 공통 표시)."""
    lines = block_text.strip().split('\n')
    title = lines[0].replace('### 📄', '').strip()
    body = "\n".join(lines[1:]).strip()

    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'

    row_title = table.add_row()
    set_cell_background(row_title.cells[0], "FFF3D6")
    add_parsed_text_to_cell(row_title.cells[0], f"📄 {title}")

    row_body = table.add_row()
    add_parsed_text_to_cell(row_body.cells[0], body if body else "내용 없음")

    doc.add_paragraph()
def add_question_block_to_doc(doc, block_text, is_teacher):
    """### 📌 로 시작하는 '문항' 블록을 워드 문서에 렌더링합니다."""
    lines = block_text.strip().split('\n')
    title = lines[0].replace('### 📌', '').strip()
    body = "\n".join(lines[1:])

    q_match = re.search(r'\[질문\](.*?)(?=\[평가의도\]|\[모범답안\]|\[꼬리질문\]|$)', body, re.DOTALL)
    i_match = re.search(r'\[평가의도\](.*?)(?=\[모범답안\]|\[꼬리질문\]|$)', body, re.DOTALL)
    a_match = re.search(r'\[모범답안\](.*?)(?=\[꼬리질문\]|$)', body, re.DOTALL)
    f_match = re.search(r'\[꼬리질문\](.*?)(?=$)', body, re.DOTALL)

    q_text = q_match.group(1).strip() if q_match else "내용 없음"
    i_text = i_match.group(1).strip() if i_match else ""
    a_text = a_match.group(1).strip() if a_match else ""
    f_text = f_match.group(1).strip() if f_match else ""

    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'

    row_title = table.add_row()
    set_cell_background(row_title.cells[0], "EBF1FA")
    add_parsed_text_to_cell(row_title.cells[0], f"📌 {title}")

    row_q = table.add_row()
    add_parsed_text_to_cell(row_q.cells[0], f"**[면접 질문]**\n{q_text}")

    if is_teacher:
        row_i = table.add_row()
        set_cell_background(row_i.cells[0], "F9F9F9")
        add_parsed_text_to_cell(row_i.cells[0], f"**[평가 의도]**\n{i_text}")

        row_a = table.add_row()
        add_parsed_text_to_cell(row_a.cells[0], f"**[모범 답안 가이드]**\n{a_text}")

        row_f = table.add_row()
        set_cell_background(row_f.cells[0], "FFF4F4")
        add_parsed_text_to_cell(row_f.cells[0], f"**[압박용 꼬리질문]**\n{f_text}")

    doc.add_paragraph()
def add_source_block_to_doc(doc, block_text):
    """### 📊 로 시작하는 '출제근거' 요약 블록(서울대 기출 문서 양식 참고)을 워드 문서에 렌더링합니다.
    실제 시험지에서 이 정보는 수험생이 아닌 출제/채점 참고 자료이므로 교사용 문서에만 표시합니다."""
    lines = block_text.strip().split('\n')
    title = lines[0].replace('### 📊', '').strip()
    body = "\n".join(lines[1:])

    u_match = re.search(r'\[활용모집단위\](.*?)(?=\[개념\]|\[교육과정출처\]|\[자료출처\]|$)', body, re.DOTALL)
    c_match = re.search(r'\[개념\](.*?)(?=\[교육과정출처\]|\[자료출처\]|$)', body, re.DOTALL)
    s_match = re.search(r'\[교육과정출처\](.*?)(?=\[자료출처\]|$)', body, re.DOTALL)
    r_match = re.search(r'\[자료출처\](.*?)(?=$)', body, re.DOTALL)

    u_text = u_match.group(1).strip() if u_match else ""
    c_text = c_match.group(1).strip() if c_match else ""
    s_text = s_match.group(1).strip() if s_match else ""
    r_text = r_match.group(1).strip() if r_match else ""

    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'

    row_title = table.add_row()
    set_cell_background(row_title.cells[0], "EAF6EC")
    add_parsed_text_to_cell(row_title.cells[0], f"📊 {title} (서울대 기출 문서 양식 참고 · 교사 참고용)")

    row_u = table.add_row()
    add_parsed_text_to_cell(row_u.cells[0], f"**[활용 모집단위]**\n{u_text}")

    row_c = table.add_row()
    set_cell_background(row_c.cells[0], "F9F9F9")
    add_parsed_text_to_cell(row_c.cells[0], f"**[핵심 개념]**\n{c_text}")

    row_s = table.add_row()
    add_parsed_text_to_cell(row_s.cells[0], f"**[교육과정 출처]**\n{s_text}")

    row_r = table.add_row()
    set_cell_background(row_r.cells[0], "F9F9F9")
    add_parsed_text_to_cell(row_r.cells[0], f"**[참고 자료출처]**\n{r_text}")

    doc.add_paragraph()
def create_word_files(content, student_name, interview_type, target_desc):
    type_label = "생기부면접" if "생기부" in interview_type else "제시문면접"

    doc_student = Document()
    doc_teacher = Document()

    # '### 📄'(제시문) · '### 📌'(문항) · '### 📊'(출제근거 요약, 서울대 기출 문서 양식 참고)
    # 마커를 기준으로 등장 순서 그대로 섹션을 나눔
    # → 제시문 면접의 '제시문 1개 + 문제 2개(+출제근거 요약)'가 교차되는 구조를 그대로 지원
    sections = re.split(r'(?=### 📄|### 📌|### 📊)', content)

    for doc, is_teacher in [(doc_student, False), (doc_teacher, True)]:
        set_document_font(doc)
        title_text = f"🎓 [{student_name}] {target_desc} 도개고 맞춤 모의면접 {'지침서 (교사용)' if is_teacher else '워크북 (학생용)'}"
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(f"[{type_label}] 본 문서는 도개고등학교 진로진학 지도 기준에 맞춰 생성되었습니다.\n")

        for section in sections:
            section = section.strip('\n')
            if not section.strip():
                continue
            if section.strip().startswith('### 📄'):
                add_passage_block_to_doc(doc, section)
            elif section.strip().startswith('### 📌'):
                add_question_block_to_doc(doc, section, is_teacher)
            elif section.strip().startswith('### 📊'):
                if is_teacher:
                    add_source_block_to_doc(doc, section)
            else:
                # 마커가 나오기 전의 서론/설명 텍스트
                add_intro_paragraphs(doc, section.strip())
                doc.add_paragraph()

    student_path = f"{student_name}_{target_desc}_{type_label}_학생용.docx"
    teacher_path = f"{student_name}_{target_desc}_{type_label}_교사용.docx"
    doc_student.save(student_path)
    doc_teacher.save(teacher_path)
    return student_path, teacher_path
def create_chat_history_word(chat_history, student_name):
    doc = Document()
    set_document_font(doc)
    doc.add_heading(f"💬 [{student_name}] 면접 문항 피드백 대화 내역", level=1)
    doc.add_paragraph("AI 출제위원과의 피드백 기록입니다.\n" + "="*50)

    for msg in chat_history:
        role_title = "👤 선생님 (요청)" if msg["role"] == "user" else "🤖 AI 출제위원 (답변/수정본)"
        doc.add_heading(role_title, level=2)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        parts = msg["content"].split("**")
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 != 0: run.bold = True
        doc.add_paragraph("-" * 50)
    file_path = f"{student_name}_피드백_대화내역.docx"
    doc.save(file_path)
    return file_path
# -------------------------------------------------------------------------
# [3] 메인 UI
# -------------------------------------------------------------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "word_files" not in st.session_state:
    st.session_state.word_files = None
if "raw_content" not in st.session_state:
    st.session_state.raw_content = ""
st.title("🎓 도개고 대입 모의면접 마스터 솔루션")
# 🌟 [업데이트] OCR 변환 가이드가 포함된 상세 설명서 🌟
with st.expander("📖 [클릭] 프로그램 사용 설명서 및 PDF OCR 변환 방법", expanded=False):
    st.markdown("""
    ### 🔑 Google Gemini API 키 발급 방법
    1. **Google AI Studio 접속:** [Google AI Studio](https://aistudio.google.com/)에 접속합니다.
    2. **구글 계정 로그인:** 평소 사용하는 구글 계정으로 로그인합니다.
    3. **API 키 생성:** 좌측 상단 'Get API key' 버튼을 클릭하여 키 생성 후 복사합니다.
    ### 📂 [중요] 생기부 PDF는 반드시 '텍스트 추출(OCR)'된 파일이어야 합니다!
    * **왜 필요한가요?** 단순 이미지(스캔본) PDF는 AI가 글자를 읽지 못하므로, **마우스로 글자가 드래그되거나 텍스트로 인식되는 PDF**여야만 정상 분석이 가능합니다.
    * **스캔된 PDF를 OCR(텍스트형)로 변환하는 방법:**
      1. **구글 드라이브(Google Drive) 활용 (가장 추천):**
         - 스캔된 생기부 PDF 파일을 구글 드라이브에 업로드합니다.
         - 업로드된 파일 우클릭 ➔ **[연결 앱]** ➔ **[Google 문서]**를 선택하여 엽니다.
         - 구글 문서로 열리면 이미지가 텍스트로 자동 변환됩니다. 상단 메뉴 **[파일] ➔ [다운로드] ➔ [Microsoft Word(.docx)]** 또는 PDF로 저장합니다.
      2. **온라인 무료 툴 활용:**
         - 'ILOVEPDF' 또는 'Smallpdf' 등의 사이트에서 **'OCR PDF(PDF 텍스트 변환)'** 메뉴를 이용해 변환합니다.
    ### 💡 프로그램 사용 순서
    1. 왼쪽 사이드바에 API 키를 입력합니다.
    2. 면접 방식(생기부 기반/제시문 기반)을 선택합니다.
    3. 대학, 학과, 난이도를 설정합니다.
    4. 생기부 면접일 경우 **OCR 변환된 PDF**를 업로드하고 생성 버튼을 누릅니다.
    """)
with st.sidebar:
    api_key = st.text_input("🔑 Gemini API Key", type="password")
UNIVERSITIES = {
    "서울권": ["서울대", "연세대", "고려대", "성균관대", "서강대", "한양대", "중앙대", "경희대", "한국외대", "서울시립대", "이화여대"],
    "충청권": ["카이스트(KAIST)", "충남대", "충북대", "고려대(세종)"],
    "경상권": ["경북대", "부산대", "UNIST", "영남대", "계명대"]
}
col1, col2 = st.columns(2)
with col1:
    interview_type = st.radio("🎯 면접 방식", ["생기부 기반 면접", "상위권 대학 제시문 기반 면접"], horizontal=True)
    region = st.selectbox("📍 권역 선택", ["서울권", "충청권", "경상권", "직접 입력"])
    if region == "직접 입력":
        uni = st.text_input("🏫 대학 직접 입력", value="한국대")
    else:
        uni = st.selectbox("🏫 대학 선택", UNIVERSITIES[region])
with col2:
    major = st.text_input("🎓 지원 학과/전공", placeholder="예: 철학과")
    student_name = st.text_input("👤 지원자 성명", value="김기섭")
    difficulty = st.radio("⚙️ 난이도 선택", ["하 (기초)", "중 (표준)", "상 (압박)"], horizontal=True, index=1)
uploaded_file = None
if interview_type == "생기부 기반 면접":
    uploaded_file = st.file_uploader("📂 학생 생기부 PDF 업로드 (OCR 변환 필수)", type=["pdf"])
st.markdown("---")
# -------------------------------------------------------------------------
# [4] 생성 버튼 로직
# -------------------------------------------------------------------------
target_desc = f"{uni}_{major}"
TEMPLATE_SANGBU = """
### 📌 [영역: OOO] **[과목명/활동명]**
[질문]
(내용 작성)
[평가의도]
(내용 작성)
[모범답안]
(내용 작성)
[꼬리질문]
(내용 작성)
"""
# 서울대학교 "2026학년도 대학 신입학생 수시모집 일반전형 면접 및 구술고사" 문서 양식을 참고한 구조.
# '제시문(가·나·(다)) 1세트 + 문제 2개 + 출제근거 요약' = 1세트, 총 3세트(제시문 3세트 · 문제 6개) 구조
TEMPLATE_JESIMUN = """
### 📄 [세트 1] **[분류: 인문학 등 학과에 맞는 계열 1개 선택]**
(가) (서울대 기출처럼 정의-원리-예시 흐름을 가진 압축적·개념적 지문 4~6문장, 학술적 문어체)
(나) (가)와 대비되거나 보완되는 별도 관점의 지문 4~6문장)
(다) (필요한 경우에만 추가. 불필요하면 이 줄 자체를 생략)
### 📌 [세트1-문제1]
[질문]
[문제 1] (가)와 (나)의 관점을 비교·분석하거나 하나의 입장에서 다른 하나를 비판하도록 요구하는 질문
[평가의도]
(이 문제로 평가하려는 분석적·비판적 사고력을 구체적으로 서술)
[모범답안]
(핵심 논리 흐름을 담은 모범 답안 가이드 + 학생이 막혔을 때 면접관이 줄 수 있는 힌트 1가지)
[꼬리질문]
(1단계: 답변을 더 이끌어내는 꼬리질문 / 2단계: 학생의 논리·가정에 대한 반론을 제기하는 심화 꼬리질문)
### 📌 [세트1-문제2]
[질문]
[문제 2] (가), (나)(, (다))의 유기적 관계를 종합하거나 새로운 상황에 적용·확장하도록 요구하는 질문
[평가의도]
(내용 작성)
[모범답안]
(내용 작성)
[꼬리질문]
(1단계 / 2단계 구조로 내용 작성)
### 📊 [세트 1 출제근거]
[활용모집단위]
(이 세트가 어울리는 모집단위 예시 2~3개, 지원 학과 포함)
[개념]
(핵심 개념 키워드 3~5개, 쉼표로 구분)
[교육과정출처]
(관련 과목·단원명 1~3개)
[자료출처]
(참고할 만한 문헌/자료 예시 1~2개)
### 📄 [세트 2] **[분류: 세트 1과 다른 계열 1개 선택]**
(가) (지문 4~6문장)
(나) (지문 4~6문장)
### 📌 [세트2-문제1]
[질문]
(내용 작성)
[평가의도]
(내용 작성)
[모범답안]
(내용 작성)
[꼬리질문]
(1단계 / 2단계 구조로 내용 작성)
### 📌 [세트2-문제2]
[질문]
(내용 작성)
[평가의도]
(내용 작성)
[모범답안]
(내용 작성)
[꼬리질문]
(1단계 / 2단계 구조로 내용 작성)
### 📊 [세트 2 출제근거]
[활용모집단위]
(내용 작성)
[개념]
(내용 작성)
[교육과정출처]
(내용 작성)
[자료출처]
(내용 작성)
### 📄 [세트 3] **[분류: 세트 1·2와 다른 계열 1개 선택]**
(가) (지문 4~6문장)
(나) (지문 4~6문장)
### 📌 [세트3-문제1]
[질문]
(내용 작성)
[평가의도]
(내용 작성)
[모범답안]
(내용 작성)
[꼬리질문]
(1단계 / 2단계 구조로 내용 작성)
### 📌 [세트3-문제2]
[질문]
(내용 작성)
[평가의도]
(내용 작성)
[모범답안]
(내용 작성)
[꼬리질문]
(1단계 / 2단계 구조로 내용 작성)
### 📊 [세트 3 출제근거]
[활용모집단위]
(내용 작성)
[개념]
(내용 작성)
[교육과정출처]
(내용 작성)
[자료출처]
(내용 작성)
"""
if st.button("🚀 면접 패키지 생성 시작"):
    if not api_key: st.error("API 키를 입력해 주세요."); st.stop()
    if not major: st.error("지원 학과를 입력해 주세요."); st.stop()
    if interview_type == "생기부 기반 면접" and not uploaded_file: st.error("생기부 파일을 업로드해 주세요."); st.stop()

    student_record = extract_text_from_pdf(uploaded_file) if uploaded_file else ""

    if interview_type == "생기부 기반 면접":
        prompt = f"""
        당신은 도개고등학교의 진학 지도 노하우와 {uni} {major} 입학사정관의 시각을 겸비한 최고급 면접 출제위원입니다.
        지원자 '{student_name}' 학생의 생기부를 분석하여 도개고 선배들이 실제 상위권/지역거점 대학 면접에서 마주했던 수준 높은 기출 문항들의 난이도와 깊이를 반영해 질문을 만드세요[cite: 3, 4].
        면접 난이도: {difficulty}

        [실제 합격 선배들의 증언 기반 출제 가이드]
        서울대학교 합격 선배들의 후기에 따르면 학생부(학교생활기록부) 기반 면접은 다음과 같은 특징을 보입니다. 이를 최대한 반영해 질문을 설계하세요.
        - 활동/동아리명 등에 들어간 핵심 용어 자체의 정의·한계·실제 활용 사례를 직접 캐묻는 질문이 자주 나옵니다. (예: '통계'라는 단어의 의미, 한계, 활용 사례를 각각 질문)
        - 탐구에 활용한 책이나 이론의 한계점, 혹은 다루지 않은 유사 사례·제도와의 비교를 요구하며 탐구의 총체성과 깊이를 점검합니다. (예: 특정 저서의 한계점, 조선시대 유사 제도와의 비교)
        - 단순히 "무엇을 했는가"가 아니라 "왜 그 방법을 사용했는가", "그 결과로 새롭게 알게 된 것은 무엇인가", "더 깊이 탐구했다면 무엇을 추가로 다뤘어야 했는가"를 묻습니다.
        - 정답 자체보다 학생이 당황스러운 질문에도 자신의 논리로 풀어내는 사고 과정과 진정성을 확인하려 합니다.

        [지시사항]
        1. 생기부 5대 영역(1. 교과세특, 2. 창체, 3. 동아리, 4. 행특, 5. 독서/기타)을 모두 뒤져서 총 5세트를 만드세요.
        2. 과목명이나 주요 활동명은 반드시 **[생활과 윤리]** 처럼 볼드체로 묶어주세요.
        3. 단순 암기식 질문이 아니라 학생부 활동의 진위 여부, 동기, 심화 탐구 과정, 실생활 적용력을 파악하는 도개고 스타일의 날카로운 꼬리질문을 포함하세요[cite: 3, 4]. 특히 위 가이드처럼 활동 속 핵심 개념·용어의 정의/한계/활용 사례를 직접 캐묻는 꼬리질문을 최소 2세트 이상에 포함하세요.

        [출력 템플릿 엄수 - 파싱을 위해 키워드 대괄호를 절대 변경하지 마세요]
        {TEMPLATE_SANGBU}

        [생기부 내용]
        {student_record}
        """
    else:
        prompt = f"""
        당신은 서울대학교를 비롯한 상위권 대학의 실제 제시문 면접(구술고사) 출제 방식을 정확히 이해하고 있는, 도개고등학교 진학 지도 기준에 맞춘 제시문 면접 출제위원입니다.
        {major} 전공적합성과 논리력을 평가하기 위한 고난도 제시문 기반 면접을 출제하세요.
        면접 난이도: {difficulty}

        [실제 서울대학교 기출 문서 형식 참고]
        아래는 실제 "2026학년도 대학 신입학생 수시모집 일반전형 면접 및 구술고사 문항" 자료의 형식입니다. 반드시 이 형식과 문체를 참고하여 출제하세요.
        - 제시문은 (가), (나)(, 필요시 (다))처럼 문자로 구분된 여러 개의 독립적 지문으로 구성되며, 각 지문은 정의-원리-예시의 흐름을 가진 압축적이고 개념적인 문단(4~6문장)입니다. 구어체나 쉬운 설명이 아니라 학술적 문어체를 사용하세요.
        - 지문 아래에는 [문제 1], [문제 2]처럼 큰 문항이 제시되고, 각 문항은 "(가)와 (나)의 관점 중 하나를 선택하여 다른 하나를 비판하시오", "(가), (나), (다)의 유기적 관계를 파악하여 이 주장을 뒷받침하시오"처럼 제시문 간의 비교·통합·적용을 요구합니다.
        - 실제 문서에는 문항마다 [활용 모집단위], [문항해설], [출제의도], [교육과정 출제근거([개념]/[출처])], [자료출처] 정보가 표로 함께 제공됩니다. 이를 참고해 각 세트 끝에 출제근거 요약을 포함하세요.

        [선배들의 실제 면접 경험 기반 출제 가이드]
        - 같은 제시문에 대해 보통 1단계 꼬리질문은 학생의 답변을 더 이끌어내는 방향이고, 2단계 꼬리질문은 학생이 가정한 상황에 대한 반론이나 논리의 허점을 짚는 심화형 질문입니다. 이 2단계 구조를 꼬리질문에 반영하세요.
        - 서울대 제시문 면접은 정답 자체보다 논리적으로 사고를 전개하는 과정을 평가합니다. 모범답안에는 정답뿐 아니라 "학생이 막혔을 때 면접관이 줄 수 있는 힌트"도 함께 제시하세요.

        [지시사항]
        1. 생기부 내용은 무시하세요. {major} 학과와 관련된 고난도 딜레마 상황이나 철학적/학술적 주제가 담긴 제시문을 위 실제 기출 형식에 맞춰 창작하세요[cite: 3, 4].
        2. **'제시문 세트(가·나 지문 포함) 1개 + 문제 2개'를 1세트로 구성**하여, 서로 다른 주제/계열(예: 인문학, 사회과학, 수학, 과학 등)의 세트 3개(총 3세트, 문제 6개)를 만드세요.
           즉 [세트 1] → [문제 1][문제 2], [세트 2] → [문제 1][문제 2], [세트 3] → [문제 1][문제 2] 구조를 반드시 지키세요.
        3. 같은 세트 안의 두 문제는 서로 다른 각도(예: 비교·비판 vs 종합·적용)에서 접근하도록 설계하세요.
        4. 각 세트 마지막에는 실제 기출 문서처럼 [활용모집단위], [개념], [교육과정출처], [자료출처]를 요약한 출제근거 블록을 반드시 포함하세요.
        5. 서론이나 인사말은 절대 쓰지 말고, 바로 '### 📄 [세트 1]' 부터 출력하세요.

        [출력 템플릿 엄수 - 파싱을 위해 키워드 대괄호를 절대 변경하지 마세요]
        {TEMPLATE_JESIMUN}
        """

    with st.spinner(f"⏳ 로딩중... 도개고 면접 기출 수준에 맞춰 {interview_type} 문항을 조립하고 있습니다."):
        try:
            result_text = call_gemini(prompt, api_key)
            st.session_state.raw_content = result_text

            stu_path, tea_path = create_word_files(result_text, student_name, interview_type, target_desc)
            st.session_state.word_files = (stu_path, tea_path)

            display_text = result_text.replace('[질문]', '\n**💡 [면접 질문]**\n').replace('[평가의도]', '\n**🎯 [평가 의도]**\n').replace('[모범답안]', '\n**✅ [모범 답안 가이드]**\n').replace('[꼬리질문]', '\n**🔥 [압박용 꼬리질문]**\n').replace('[활용모집단위]', '\n**🏫 [활용 모집단위]**\n').replace('[개념]', '\n**🔑 [핵심 개념]**\n').replace('[교육과정출처]', '\n**📚 [교육과정 출처]**\n').replace('[자료출처]', '\n**📖 [참고 자료출처]**\n')

            # 대화 유도 문구 추가
            full_display_text = display_text + "\n\n---\n💬 **방금까지 나눈 문항 내용과 피드백 대화 내용을 한글 문서(.docx)로 만들어 드릴까요?** (원하시면 **'그래 만들어줘'**라고 말씀해 주세요!)\n\n➕ 문항이 더 필요하시면 **'더 만들어줘'** 처럼 요청해 주세요. 요청 시 기존 내용은 그대로 유지한 채 **최소 2세트가 추가**됩니다."

            st.session_state.chat_history = [{"role": "assistant", "content": full_display_text}]
            st.success("🎉 도개고 맞춤형 면접 패키지 및 워드 문서 생성이 완료되었습니다!")

        except Exception as e:
            st.error(f"❌ 생성 실패: {e}")
# -------------------------------------------------------------------------
# [5] 결과 대시보드
# -------------------------------------------------------------------------
if st.session_state.chat_history:
    st.markdown("## 📋 면접 문항 대시보드 및 실시간 피드백")

    if st.session_state.word_files:
        stu_path, tea_path = st.session_state.word_files
        chat_path = create_chat_history_word(st.session_state.chat_history, student_name)

        col_w1, col_w2, col_w3 = st.columns(3)
        with col_w1:
            with open(stu_path, "rb") as f:
                st.download_button("📥 학생용 워크북 (.docx)", f, file_name=stu_path, use_container_width=True)
        with col_w2:
            with open(tea_path, "rb") as f:
                st.download_button("📥 교사용 지침서 (.docx)", f, file_name=tea_path, use_container_width=True)
        with col_w3:
            with open(chat_path, "rb") as f:
                st.download_button("💬 피드백 대화 내역 (.docx)", f, file_name=chat_path, use_container_width=True)

    st.divider()

    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if user_feedback := st.chat_input("질문을 더 어렵게 하거나 특정 활동을 추가해 달라고 피드백을 남겨보세요. ('더 만들어줘'라고 하면 세트가 추가됩니다)"):
        st.session_state.chat_history.append({"role": "user", "content": user_feedback})
        with st.chat_message("user"):
            st.markdown(user_feedback)

        with st.chat_message("assistant"):
            with st.spinner("요청하신 피드백을 반영하여 처리 중입니다..."):

                # "더 만들어줘" 류의 '추가 생성' 요청인지 먼저 판별 (문서 저장 요청과 겹치지 않도록 우선 체크)
                add_more_words = ["더 만들어", "더만들어", "추가로 만들어", "추가해줘", "더 추가",
                                   "세트 추가", "더 뽑아", "하나 더", "한 세트 더", "두 세트 더",
                                   "더 줘", "더 필요", "더 출제"]
                is_add_more_request = any(w in user_feedback for w in add_more_words)

                doc_request_words = ["만들어", "생성", "다운", "파일", "문서로", "저장", "그래", "응", "네", "해줘"]
                is_doc_request = (not is_add_more_request) and any(w in user_feedback for w in doc_request_words) and len(user_feedback.strip()) < 15

                if is_doc_request:
                    chat_path = create_chat_history_word(st.session_state.chat_history, student_name)
                    response_text = "네! 지금까지 나눈 대화 내용을 깔끔한 워드 문서로 생성했습니다. 상단 또는 아래의 **'💬 피드백 대화 내역 (.docx)'** 다운로드 버튼을 클릭해 주세요!"
                    st.markdown(response_text)
                    st.session_state.chat_history.append({"role": "assistant", "content": response_text})
                    st.rerun()
                else:
                    required_template = TEMPLATE_SANGBU if interview_type == "생기부 기반 면접" else TEMPLATE_JESIMUN
                    previous_content = st.session_state.get("raw_content", "")

                    if is_add_more_request:
                        if interview_type == "생기부 기반 면접":
                            add_more_instruction = f"""
                            아래는 지금까지 생성된 기존 문항 전체입니다. 기존 문항은 절대 삭제·수정하지 말고 그대로 유지한 채,
                            사용자의 요청 "{user_feedback}"에 맞춰 **새로운 문항을 최소 2세트 추가**로 작성하세요.
                            (기존에 다루지 않은 생기부 영역/활동을 우선적으로 선택해 중복을 피하세요.)
                            최종 결과물에는 기존 문항 전체 + 새로 추가한 최소 2세트가 모두 포함되어야 합니다.

                            [기존 문항 전체]
                            {previous_content}
                            """
                        else:
                            add_more_instruction = f"""
                            아래는 지금까지 생성된 기존 제시문 및 문항 전체입니다. 기존 내용은 절대 삭제·수정하지 말고 그대로 유지한 채,
                            사용자의 요청 "{user_feedback}"에 맞춰 **새로운 제시문 세트를 최소 2세트 추가**로 작성하세요.
                            (새로 추가하는 세트도 반드시 '(가)·(나) 제시문 + 문제 2개 + 출제근거 요약 = 1세트' 구조를 지키고,
                            기존 세트와 겹치지 않는 새로운 주제/계열로 작성하세요. 세트 번호는 기존 마지막 번호 다음부터 이어서 매기세요.)
                            최종 결과물에는 기존 제시문·문항 전체 + 새로 추가한 최소 2세트(제시문 세트 2개, 문제 4개 이상)가 모두 포함되어야 합니다.

                            [기존 제시문 및 문항 전체]
                            {previous_content}
                            """
                        feedback_prompt = f"""
                        당신은 도개고 맞춤형 면접 출제위원입니다. 도개고 기출 수준의 심도 있는 학술적/실천적 깊이를 유지하면서
                        **반드시 다음 템플릿 구조와 [키워드]를 토씨 하나 틀리지 말고 유지**해 주세요[cite: 3, 4].

                        [강제 유지 템플릿]
                        {required_template}

                        {add_more_instruction}

                        최종 출력은 기존 내용 + 새로 추가된 내용을 합친 전체 결과물이어야 합니다. 서론·안내 문구 없이 바로 문항 내용부터 출력하세요.
                        """
                    else:
                        feedback_prompt = f"""
                        당신은 도개고 맞춤형 면접 출제위원입니다. 아래 [기존 문항 전체]를 사용자의 피드백에 맞게 수정하되,
                        도개고 기출 수준의 심도 있는 학술적/실천적 깊이를 유지하면서 **반드시 다음 템플릿 구조와 [키워드]를 토씨 하나 틀리지 말고 유지**해 주세요[cite: 3, 4].

                        [강제 유지 템플릿]
                        {required_template}

                        [기존 문항 전체]
                        {previous_content}

                        사용자 피드백: "{user_feedback}"

                        최종 출력은 피드백이 반영된 전체 결과물이어야 합니다. 서론·안내 문구 없이 바로 문항 내용부터 출력하세요.
                        """
                    try:
                        new_result = call_gemini(feedback_prompt, api_key)
                        display_text = new_result.replace('[질문]', '\n**💡 [면접 질문]**\n').replace('[평가의도]', '\n**🎯 [평가 의도]**\n').replace('[모범답안]', '\n**✅ [모범 답안 가이드]**\n').replace('[꼬리질문]', '\n**🔥 [압박용 꼬리질문]**\n').replace('[활용모집단위]', '\n**🏫 [활용 모집단위]**\n').replace('[개념]', '\n**🔑 [핵심 개념]**\n').replace('[교육과정출처]', '\n**📚 [교육과정 출처]**\n').replace('[자료출처]', '\n**📖 [참고 자료출처]**\n')

                        full_response = display_text + "\n\n---\n💬 **방금까지 나눈 문항 내용과 피드백 대화 내용을 한글 문서(.docx)로 만들어 드릴까요?** (원하시면 **'그래 만들어줘'**라고 말씀해 주세요!)\n\n➕ 문항이 더 필요하시면 **'더 만들어줘'** 처럼 요청해 주세요. 요청 시 기존 내용은 그대로 유지한 채 **최소 2세트가 추가**됩니다."

                        st.markdown(full_response)
                        st.session_state.chat_history.append({"role": "assistant", "content": full_response})

                        st.session_state.raw_content = new_result
                        stu_path, tea_path = create_word_files(new_result, student_name, interview_type, target_desc)
                        st.session_state.word_files = (stu_path, tea_path)
                        st.rerun()

                    except Exception as e:
                        st.error(f"피드백 반영 중 오류가 발생했습니다: {e}")
